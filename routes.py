from flask import render_template, request, redirect, url_for, flash, session, abort, Response, jsonify
from config import create_app, db
from models import User, HR, Candidate, Test, TestQuestion, TestAssignment, TestResult, Employee, OfficeMap, MapAccess, Report, HRNote
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from datetime import datetime, timezone
import os
import base64
from sqlalchemy import or_
from utils import is_valid_email, is_valid_phone, save_uploaded_file
    
app = create_app()

@app.route('/')
def entry():
    return render_template('entry.html')

@app.route('/admin')
def admin():
    if 'user_id' not in session or session.get('user_role') != 'admin':
        flash('Доступ запрещен. Требуются права администратора', 'error')
        return redirect(url_for('login'))
    
    try:
        users = User.query.filter(User.mail != app.config['ADMIN_EMAIL']).all()
        
        hr_list = db.session.query(
            HR.id.label('hr_id'),
            User.id.label('user_id'),
            User.surname,
            User.name,
            User.phone,
            User.mail,
            db.func.count(Candidate.id).label('candidates_count')
        ).join(
            User, HR.user_id == User.id
        ).filter(
            User.mail != app.config['ADMIN_EMAIL']
        ).outerjoin(
            Candidate, HR.id == Candidate.hr_id
        ).group_by(
            HR.id, User.id, User.surname, User.name, User.phone, User.mail
        ).all()
        
        candidates_data = db.session.query(
            Candidate.id.label('candidate_id'),
            Candidate.hr_id,
            Candidate.status,
            User.id.label('user_id'),
            User.surname,
            User.name,
            User.phone,
            User.mail
        ).join(
            User, Candidate.user_id == User.id
        ).all()
        
        hr_candidates = {}
        for hr in hr_list:
            hr_candidates[hr.hr_id] = {
                'hr_info': {
                    'id': hr.hr_id,
                    'user_id': hr.user_id,
                    'surname': hr.surname,
                    'name': hr.name,
                    'phone': hr.phone,
                    'mail': hr.mail,
                    'candidates_count': hr.candidates_count
                },
                'candidates': []
            }
        
        for candidate in candidates_data:
            if candidate.hr_id in hr_candidates:
                hr_candidates[candidate.hr_id]['candidates'].append({
                    'id': candidate.candidate_id,
                    'user_id': candidate.user_id,
                    'surname': candidate.surname,
                    'name': candidate.name,
                    'phone': candidate.phone,
                    'mail': candidate.mail,
                    'status': candidate.status
                })
        
        tests_resumes = db.session.query(
            User.id.label('user_id'),
            User.surname,
            User.name,
            Candidate.id.label('candidate_id'),
            Candidate.resume_filename,
            Candidate.application_date,
            TestAssignment.id.label('assignment_id'),
            TestAssignment.completed,
            Test.name.label('test_name'),
            TestResult.score,
            TestResult.max_score
        ).join(
            Candidate, User.id == Candidate.user_id
        ).outerjoin(
            TestAssignment, Candidate.id == TestAssignment.candidate_id
        ).outerjoin(
            Test, TestAssignment.test_id == Test.id
        ).outerjoin(
            TestResult, TestAssignment.id == TestResult.assignment_id
        ).filter(
            User.mail != app.config['ADMIN_EMAIL']
        ).all()
        
        employees = User.query.filter(
            User.role == 'employee',
            User.mail != app.config['ADMIN_EMAIL']
        ).all()
        
        stats = {
            'total_users': len(users),
            'candidates': len([u for u in users if u.role == 'candidate']),
            'hr_managers': len(hr_list),
            'employees': len(employees)
        }
        
        return render_template('admin.html',
                            current_user={
                                'name': session.get('user_name', 'Администратор'),
                                'role': 'admin'
                            },
                            users=users,
                            hr_candidates=hr_candidates,
                            employees=employees,
                            tests_resumes=tests_resumes,
                            stats=stats,
                            current_year=datetime.now().year)
    
    except Exception as e:
        flash(f'Произошла ошибка: {str(e)}', 'error')
        return redirect(url_for('admin'))
    
@app.route('/admin/change_password', methods=['POST'])
def admin_change_password():
    if 'user_id' not in session or session.get('user_role') != 'admin':
        return jsonify({'error': 'Unauthorized'}), 401
    
    data = request.get_json()
    user_id = data.get('user_id')
    new_password = data.get('new_password')
    
    if not user_id or not new_password:
        return jsonify({'error': 'Missing user_id or new_password'}), 400
    
    if len(new_password) < 5:
        return jsonify({'error': 'Password must be at least 5 characters'}), 400
    
    try:
        user = User.query.get(user_id)
        if not user:
            return jsonify({'error': 'User not found'}), 404
        
        user.password = generate_password_hash(new_password)
        db.session.commit()
        
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/admin/check_password', methods=['POST'])
def admin_check_password():
    if 'user_id' not in session or session.get('user_role') != 'admin':
        return jsonify({'error': 'Unauthorized'}), 401
    
    data = request.get_json()
    user_id = data.get('user_id')
    
    if not user_id:
        return jsonify({'error': 'Missing user_id'}), 400
    
    try:
        user = User.query.get(user_id)
        if not user:
            return jsonify({'error': 'User not found'}), 404
        
        return jsonify({
            'has_password': bool(user.password),
            'password_strength': 'strong' if len(user.password) > 50 else 'weak'
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    
@app.route('/admin/delete_user/<int:user_id>', methods=['POST'])
def admin_delete_user(user_id):
    if 'user_id' not in session or session.get('user_role') != 'admin':
        return jsonify({'error': 'Unauthorized'}), 401
    
    try:
        user = User.query.get(user_id)
        if not user:
            return jsonify({'error': 'User not found'}), 404
        
        if user.role == 'admin':
            return jsonify({'error': 'Cannot delete admin user'}), 403
        
        if user.role == 'hr':
            hr = HR.query.filter_by(user_id=user.id).first()
            if hr:
                TestAssignment.query.filter_by(assigned_by=hr.id).delete()
                OfficeMap.query.filter_by(hr_id=hr.id).delete()
                Report.query.filter_by(hr_id=hr.id).delete()
                db.session.delete(hr)
        
        elif user.role == 'candidate':
            candidate = Candidate.query.filter_by(user_id=user.id).first()
            if candidate:
                assignments = TestAssignment.query.filter_by(candidate_id=candidate.id).all()
                for assignment in assignments:
                    TestResult.query.filter_by(assignment_id=assignment.id).delete()
                    db.session.delete(assignment)
                db.session.delete(candidate)
        
        elif user.role == 'employee':
            employee = Employee.query.filter_by(user_id=user.id).first()
            if employee:
                db.session.delete(employee)
        
        db.session.delete(user)
        db.session.commit()
        
        return jsonify({'success': True})
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/admin/search_hr', methods=['GET'])
def search_hr():
    if 'user_id' not in session or session.get('user_role') != 'admin':
        return jsonify({'error': 'Unauthorized'}), 401
    
    search_term = request.args.get('term', '')
    
    try:
        hr_list = db.session.query(
            HR.id,
            User.surname,
            User.name,
            User.phone
        ).join(
            User, HR.user_id == User.id
        ).filter(
            or_(
                User.phone.ilike(f'%{search_term}%'),
                User.surname.ilike(f'%{search_term}%'),
                User.name.ilike(f'%{search_term}%')
            )
        ).limit(10).all()
        
        results = [{
            'id': hr.id,
            'name': f"{hr.surname} {hr.name}",
            'phone': hr.phone
        } for hr in hr_list]
        
        return jsonify(results)
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/admin/reassign_candidate', methods=['POST'])
def reassign_candidate():
    if 'user_id' not in session or session.get('user_role') != 'admin':
        return jsonify({'error': 'Unauthorized'}), 401
    
    data = request.get_json()
    candidate_id = data.get('candidate_id')
    new_hr_id = data.get('new_hr_id')
    
    if not candidate_id or not new_hr_id:
        return jsonify({'error': 'Missing candidate_id or new_hr_id'}), 400
    
    try:
        candidate = Candidate.query.get(candidate_id)
        if not candidate:
            return jsonify({'error': 'Candidate not found'}), 404
        
        new_hr = HR.query.get(new_hr_id)
        if not new_hr:
            return jsonify({'error': 'HR not found'}), 404
        
        candidate.hr_id = new_hr.id
        db.session.commit()
        
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/admin/get_all_data', methods=['GET'])
def admin_get_all_data():
    if 'user_id' not in session or session.get('user_role') != 'admin':
        return jsonify({'error': 'Unauthorized'}), 401
    
    try:
        users = User.query.all()
        users_data = [{
            'id': u.id,
            'name': f"{u.surname} {u.name}",
            'role': u.role,
            'email': u.mail,
            'phone': u.phone
        } for u in users]
        
        hr_list = HR.query.options(
            db.joinedload(HR.user),
            db.joinedload(HR.candidates).joinedload(Candidate.user)
        ).all()
        
        hr_candidates_data = []
        for hr in hr_list:
            for candidate in hr.candidates:
                hr_candidates_data.append({
                    'hr_id': hr.user.id,
                    'hr_name': f"{hr.user.surname} {hr.user.name}",
                    'hr_phone': hr.user.phone,
                    'candidate_id': candidate.user.id,
                    'candidate_name': f"{candidate.user.surname} {candidate.user.name}",
                    'status': candidate.status
                })
        
        candidates = Candidate.query.options(
            db.joinedload(Candidate.user),
            db.joinedload(Candidate.test_assignments).joinedload(TestAssignment.test),
            db.joinedload(Candidate.test_assignments).joinedload(TestAssignment.results)
        ).all()
        
        tests_resumes_data = []
        for candidate in candidates:
            tests_info = []
            for assignment in candidate.test_assignments:
                tests_info.append({
                    'test_name': assignment.test.name,
                    'status': 'Пройден' if assignment.completed else 'Ожидает',
                    'result': f"{assignment.results.score}/{assignment.results.max_score}" if assignment.results else '-'
                })
            
            tests_resumes_data.append({
                'candidate_id': candidate.user.id,
                'candidate_name': f"{candidate.user.surname} {candidate.user.name}",
                'tests': tests_info,
                'resume': bool(candidate.resume_filename),
                'resume_date': candidate.application_date.strftime('%d.%m.%Y') if candidate.application_date else '-'
            })
        
        return jsonify({
            'users': users_data,
            'hr_candidates': hr_candidates_data,
            'tests_resumes': tests_resumes_data,
            'stats': {
                'total_users': len(users_data),
                'candidates': len([u for u in users_data if u['role'] == 'candidate']),
                'hr_managers': len([u for u in users_data if u['role'] == 'hr']),
                'employees': len([u for u in users_data if u['role'] == 'employee'])
            }
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    
@app.route('/hire_candidate/<int:candidate_id>', methods=['POST'])
def hire_candidate(candidate_id):
    if 'user_id' not in session or session.get('user_role') != 'hr':
        return jsonify({'error': 'Unauthorized'}), 401
    
    try:
        evaluation_data = request.get_json()
        
        candidate = Candidate.query.filter_by(user_id=candidate_id).first()
        if not candidate:
            return jsonify({'error': 'Candidate not found'}), 404
        
        user = User.query.get(candidate_id)
        if not user or user.role != 'candidate':
            return jsonify({'error': 'User not found or not a candidate'}), 404
        
        evaluation_text = (
            f"Должность: {evaluation_data['position']}\n"
            f"Креативность: {evaluation_data['creativity']}/10\n"
            f"Технические навыки: {evaluation_data['tech_skills']}/10\n"
            f"Коммуникация: {evaluation_data['communication']}/10\n"
            f"Тайм-менеджмент: {evaluation_data['time_management']}/10\n"
            f"Работа в команде: {evaluation_data['teamwork']}/10"
        )
        
        user.role = 'employee'
        
        new_employee = Employee(
            user_id=user.id,
            position=evaluation_data['position'],
            additional_info=evaluation_text,
            hire_date=datetime.now(timezone.utc)
        )
        db.session.add(new_employee)
        
        db.session.delete(candidate)
        
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        print(f"Error hiring candidate: {str(e)}")
        return jsonify({'error': str(e)}), 500
    
@app.route('/create_test', methods=['GET', 'POST'])
def create_test():
    if 'user_id' not in session or session.get('user_role') != 'hr':
        flash('Доступ запрещен', 'error')
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        num_questions = int(request.form.get('num_questions', 1))
        test = Test(
            hr_id=HR.query.filter_by(user_id=session['user_id']).first().id,
            name="Новый тест",
            description=""
        )
        db.session.add(test)
        db.session.flush()
        
        for i in range(1, num_questions + 1):
            db.session.add(TestQuestion(
                test_id=test.id,
                question_text=f"Вопрос {i}",
                question_type='text',
                points=1,
                order_num=i
            ))
        
        db.session.commit()
        return redirect(url_for('edit_test', test_id=test.id))
    
    return render_template('create_test.html')

@app.route('/hr/test_results/<int:result_id>', methods=['GET', 'POST'])
def hr_test_result(result_id):
    if 'user_id' not in session or session.get('user_role') != 'hr':
        flash('Доступ запрещен', 'error')
        return redirect(url_for('login'))
    
    result = TestResult.query.get_or_404(result_id)
    test = Test.query.get_or_404(result.test_id)
    
    if result.detailed_results is None:
        result.detailed_results = []
        db.session.commit()
    
    if request.method == 'POST':
        try:
            new_scores = request.get_json()
            total_score = 0
            detailed_results = []
            
            for question in test.questions:
                question_score = new_scores.get(str(question.id), 0)
                total_score += min(question_score, question.points)
                detailed_results.append({
                    'question_id': question.id,
                    'score': question_score,
                    'max_score': question.points
                })
            
            result.score = total_score
            result.detailed_results = detailed_results
            db.session.commit()
            
            return jsonify({'success': True})
        
        except Exception as e:
            db.session.rollback()
            return jsonify({'error': str(e)}), 500
    
    return render_template('hr_test_result.html',
                         result=result,
                         test=test,
                         current_user={
                             'name': session.get('user_name')
                         })


@app.route('/edit_test', methods=['GET', 'POST'])
def edit_test():
    if 'user_id' not in session or session.get('user_role') != 'hr':
        flash('Доступ запрещен', 'error')
        return redirect(url_for('login'))
    
    hr_record = HR.query.filter_by(user_id=session['user_id']).first()
    if not hr_record:
        flash('HR запись не найдена', 'error')
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        try:
            test_data = request.get_json()
            test = Test.query.get(test_data['test_id'])
            
            if not test or test.hr_id != hr_record.id:
                return jsonify({'error': 'Тест не найден или нет прав на редактирование'}), 403
            
            test.name = test_data['test_name']
            test.description = test_data.get('test_description', '')
            
            TestQuestion.query.filter_by(test_id=test.id).delete()
            
            for i, question in enumerate(test_data['questions'], start=1):
                new_question = TestQuestion(
                    test_id=test.id,
                    question_text=question['text'],
                    question_type=question['type'],
                    options=question.get('options'),
                    correct_answer=question.get('correct_answer'),
                    points=question.get('points', 1),
                    order_num=i
                )
                db.session.add(new_question)
            
            db.session.commit()
            return jsonify({'success': True, 'test_id': test.id})
        
        except Exception as e:
            db.session.rollback()
            return jsonify({'error': str(e)}), 500
    
    test_id = request.args.get('test_id')
    if not test_id:
        return redirect(url_for('create_test'))
    
    test = Test.query.get(test_id)
    if not test or test.hr_id != hr_record.id:
        flash('Тест не найден или нет прав доступа', 'error')
        return redirect(url_for('hr'))
    
    questions = TestQuestion.query.filter_by(test_id=test.id).order_by(TestQuestion.order_num).all()
    
    return render_template('edit_test.html', 
                         test=test,
                         questions=questions)


@app.route('/assign_test', methods=['POST'])
def assign_test():
    if 'user_id' not in session or session.get('user_role') != 'hr':
        return jsonify({'error': 'Unauthorized'}), 401
    
    data = request.get_json()
    hr_user = User.query.get(session['user_id'])
    hr_record = HR.query.filter_by(user_id=hr_user.id).first()
    
    candidate = Candidate.query.filter_by(
        id=data['candidate_id'],
        hr_id=hr_record.id
    ).first()
    
    if not candidate:
        return jsonify({'error': 'Кандидат не найден'}), 404
    
    assignment = TestAssignment(
        test_id=data['test_id'],
        candidate_id=data['candidate_id'],
        assigned_by=hr_record.id
    )
    db.session.add(assignment)
    db.session.commit()
    
    return jsonify({'success': True})
            
@app.route('/reject_candidate/<int:candidate_id>', methods=['POST'])
def reject_candidate(candidate_id):
    if 'user_id' not in session or session.get('user_role') != 'hr':
        return jsonify({'error': 'Unauthorized'}), 401
    
    try:
        candidate = Candidate.query.filter_by(user_id=candidate_id).first()
        if not candidate:
            return jsonify({'error': 'Candidate not found'}), 404
        
        user = User.query.get(candidate_id)
        if user and user.role == 'candidate':
            db.session.delete(user)
        
        db.session.delete(candidate)
        
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/hr')
def hr():
    if 'user_id' not in session or session.get('user_role') != 'hr':
        flash('Доступ запрещен', 'error')
        return redirect(url_for('login'))
    
    user = db.session.get(User, session['user_id'])
    hr_record = HR.query.filter_by(user_id=user.id).first()
    
    if not hr_record:
        flash('HR запись не найдена', 'error')
        return redirect(url_for('login'))
    
    candidates = Candidate.query.filter_by(hr_id=hr_record.id).options(
        db.joinedload(Candidate.user)
    ).all()
    
    tests = Test.query.filter_by(hr_id=hr_record.id).all()
    
    assignments = TestAssignment.query.filter(
        TestAssignment.assigned_by == hr_record.id
    ).options(
        db.joinedload(TestAssignment.test),
        db.joinedload(TestAssignment.candidate).joinedload(Candidate.user),
        db.joinedload(TestAssignment.results)
    ).all()
    
    reports = Report.query.filter_by(hr_id=hr_record.id).all()
    
    return render_template('hr.html',
                         current_user={
                             'name': f"{user.surname} {user.name}",
                             'role': 'hr'
                         },
                         candidates=candidates,
                         tests=tests,
                         assignments=assignments,
                         reports=reports,
                         current_year=datetime.now().year)

@app.route('/download_candidate_resume/<int:candidate_id>')
def download_candidate_resume(candidate_id):
    if 'user_id' not in session or session.get('user_role') != 'hr':
        abort(403)
    
    candidate = Candidate.query.get_or_404(candidate_id)
    hr = HR.query.filter_by(user_id=session['user_id']).first()
    
    if not hr or candidate.hr_id != hr.id:
        abort(403)
    
    if not candidate.resume_file:
        abort(404, description="Резюме не найдено")
    
    return Response(
        candidate.resume_file,
        mimetype=candidate.resume_content_type,
        headers={
            'Content-Disposition': f'attachment; filename="{candidate.resume_filename}"'
        }
    )

@app.route('/hr/edit_test_result/<int:result_id>', methods=['GET', 'POST'])
def edit_test_result(result_id):
    if 'user_id' not in session or session.get('user_role') != 'hr':
        flash('Доступ запрещен', 'error')
        return redirect(url_for('login'))
    
    result = TestResult.query.get_or_404(result_id)
    test = Test.query.get_or_404(result.test_id)
    
    if request.method == 'POST':
        new_scores = request.get_json()
        total_score = 0
        
        for question in test.questions:
            question_score = new_scores.get(str(question.id), 0)
            total_score += min(question_score, question.points)
            
            for dr in result.detailed_results:
                if dr['question_id'] == question.id:
                    dr['score'] = question_score
        
        result.score = total_score
        db.session.commit()
        
        return jsonify({'success': True})
    
    return render_template('edit_test_result.html',
                         result=result,
                         test=test,
                         current_user={
                             'name': session.get('user_name')
                         })

@app.route('/maps', methods=['GET'])
def maps():
    if 'user_id' not in session:
        flash('Требуется авторизация', 'error')
        return redirect(url_for('login'))
    
    user = db.session.get(User, session['user_id'])
    
    if session.get('user_role') == 'hr':
        hr = HR.query.filter_by(user_id=user.id).first()
        maps = OfficeMap.query.filter_by(hr_id=hr.id).all() if hr else []
    else:
        maps = OfficeMap.query.join(MapAccess).filter(
            MapAccess.user_id == user.id
        ).all()
    
    users = User.query.all() if session.get('user_role') == 'hr' else []
    
    return render_template('maps.html',
                        current_user={
                            'name': f"{user.surname} {user.name}",
                            'role': session.get('user_role')
                        },
                        maps=maps,
                        users=users)

@app.route('/create_map', methods=['GET', 'POST'])
def create_map():
    if 'user_id' not in session or session.get('user_role') != 'hr':
        flash('Доступ запрещен. Требуются права HR', 'error')
        return redirect(url_for('login'))
    
    if request.method == 'GET':
        return render_template('map-editor.html', map=None)
    
    data = request.get_json()
    if not data or not data.get('name') or not data.get('image_data'):
        return jsonify({'error': 'Необходимо указать название и изображение'}), 400
    
    hr = HR.query.filter_by(user_id=session['user_id']).first()
    if not hr:
        return jsonify({'error': 'HR не найден'}), 404
    
    try:
        new_map = OfficeMap(
            hr_id=hr.id,
            name=data['name'],
            image_data=data['image_data']
        )
        db.session.add(new_map)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'map_id': new_map.id
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/edit_map/<int:map_id>', methods=['GET', 'POST'])
def edit_map(map_id):
    if 'user_id' not in session or session.get('user_role') != 'hr':
        flash('Доступ запрещен. Требуются права HR', 'error')
        return redirect(url_for('login'))
    
    hr = HR.query.filter_by(user_id=session['user_id']).first()
    map_data = OfficeMap.query.get_or_404(map_id)
    
    if not hr or map_data.hr_id != hr.id:
        flash('Нет прав на редактирование этой карты', 'error')
        return redirect(url_for('maps'))
    
    if request.method == 'POST':
        map_data.name = request.form.get('name', map_data.name)
        image = request.files.get('image')
        
        if image and image.filename != '':
            map_data.image_data = f"data:{image.mimetype};base64,{base64.b64encode(image.read()).decode('utf-8')}"
        
        db.session.commit()
        flash('Карта успешно обновлена', 'success')
        return redirect(url_for('maps'))
    
    return render_template('map-editor.html', map=map_data)

@app.route('/delete_map/<int:map_id>', methods=['POST'])
def delete_map(map_id):
    if 'user_id' not in session or session.get('user_role') != 'hr':
        return jsonify({'error': 'Unauthorized'}), 403
    
    hr = HR.query.filter_by(user_id=session['user_id']).first()
    map_data = OfficeMap.query.get_or_404(map_id)
    
    if not hr or map_data.hr_id != hr.id:
        return jsonify({'error': 'Forbidden'}), 403
    
    try:
        MapAccess.query.filter_by(map_id=map_id).delete()
        db.session.delete(map_data)
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/map_access/<int:map_id>', methods=['GET', 'POST'])
def map_access(map_id):
    if 'user_id' not in session or session.get('user_role') != 'hr':
        return jsonify({'error': 'Unauthorized'}), 403
    
    if request.method == 'GET':
        employees = db.session.query(Employee, User).join(
            User, Employee.user_id == User.id
        ).all()
        
        current_access = MapAccess.query.filter_by(map_id=map_id).all()
        current_user_ids = [access.user_id for access in current_access]
        
        return jsonify({
            'employees': [
                {
                    'id': user.id,
                    'name': f"{user.surname} {user.name}",
                    'position': employee.position,
                    'has_access': user.id in current_user_ids
                } for employee, user in employees
            ]
        })
    
    if request.method == 'POST':
        try:
            data = request.get_json()
            MapAccess.query.filter_by(map_id=map_id).delete()
            
            for user_id in data.get('users', []):
                access = MapAccess(map_id=map_id, user_id=user_id)
                db.session.add(access)
            
            db.session.commit()
            return jsonify({'success': True})
        except Exception as e:
            db.session.rollback()
            return jsonify({'error': str(e)}), 500
        
@app.route('/get_map_access/<int:map_id>')
def get_map_access(map_id):
    if 'user_id' not in session or session.get('user_role') != 'hr':
        return jsonify({'error': 'Unauthorized'}), 403
    
    current_access = MapAccess.query.filter_by(map_id=map_id).all()
    return jsonify({'user_ids': [access.user_id for access in current_access]})      
        
@app.route('/view_map/<int:map_id>')
def view_map(map_id):
    if 'user_id' not in session:
        flash('Требуется авторизация', 'error')
        return redirect(url_for('login'))
    
    map_data = OfficeMap.query.get_or_404(map_id)
    user = db.session.get(User, session['user_id'])
    
    if session.get('user_role') != 'hr' or map_data.hr_id != user.id:
        access = MapAccess.query.filter_by(
            map_id=map_id,
            user_id=user.id
        ).first()
        if not access:
            flash('Нет доступа к этой карте', 'error')
            return redirect(url_for('maps'))
    
    return render_template('view-map.html', map=map_data)

@app.route('/update_map/<int:map_id>', methods=['POST'])
def update_map(map_id):
    if 'user_id' not in session or session.get('user_role') != 'hr':
        return jsonify({'error': 'Unauthorized'}), 403
    
    data = request.get_json()
    if not data or not data.get('name') or not data.get('image_data'):
        return jsonify({'error': 'Необходимо указать название и изображение'}), 400
    
    hr = HR.query.filter_by(user_id=session['user_id']).first()
    map_data = OfficeMap.query.get_or_404(map_id)
    
    if not hr or map_data.hr_id != hr.id:
        return jsonify({'error': 'Нет прав на редактирование этой карты'}), 403
    
    try:
        map_data.name = data['name']
        map_data.image_data = data['image_data']
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500
    
@app.route('/get_maps')
def get_maps():
    if 'user_id' not in session or session.get('user_role') != 'hr':
        return jsonify({'error': 'Unauthorized'}), 401
    
    hr = HR.query.filter_by(user_id=session['user_id']).first()
    if not hr:
        return jsonify({'error': 'HR not found'}), 404
    
    maps = OfficeMap.query.filter_by(hr_id=hr.id).all()
    return jsonify({
        'maps': [{
            'id': map.id,
            'name': map.name,
            'image_data': map.image_data,
            'created_at': map.created_at.isoformat()
        } for map in maps]
    })

@app.route('/employee')
def employee():
    if 'user_id' not in session or session.get('user_role') != 'employee':
        flash('Доступ запрещен. Требуется авторизация как сотрудник', 'error')
        return redirect(url_for('login'))
    
    user = db.session.get(User, session['user_id'])
    employee_info = Employee.query.filter_by(user_id=user.id).first()
    
    maps = OfficeMap.query.join(MapAccess).filter(
        MapAccess.user_id == user.id
    ).all()
    
    return render_template('employee.html',
                         current_user=user,
                         employee_info=employee_info,
                         maps=maps,
                         current_year=datetime.now().year)

@app.route('/get_employees')
def get_employees():
    if 'user_id' not in session or session.get('user_role') != 'hr':
        return jsonify({'error': 'Unauthorized'}), 403
    
    employees = db.session.query(Employee, User).join(
        User, Employee.user_id == User.id
    ).all()
    
    return jsonify({
        'employees': [
            {
                'id': user.id,
                'name': f"{user.surname} {user.name}",
                'position': employee.position
            } for employee, user in employees
        ]
    })

@app.route('/take_test/<int:assignment_id>', methods=['GET', 'POST'])
def take_test(assignment_id):
    if 'user_id' not in session or session.get('user_role') != 'candidate':
        return jsonify({'error': 'Unauthorized'}), 401
    
    user = db.session.get(User, session['user_id'])
    if not user:
        flash('Пользователь не найден', 'error')
        return redirect(url_for('login'))
    
    assignment = TestAssignment.query.get_or_404(assignment_id)
    test = Test.query.get_or_404(assignment.test_id)
    
    if request.method == 'POST':
        try:
            answers = request.get_json()
            if not answers:
                return jsonify({'error': 'No answers provided'}), 400
            
            score = 0
            max_score = sum(q.points for q in test.questions)
            detailed_results = []
            
            for question in test.questions:
                user_answer = answers.get(str(question.id))
                question_score = 0
                
                if question.question_type == 'text':
                    question_score = question.points
                elif user_answer and question.correct_answer:
                    if question.question_type == 'single':
                        if user_answer in question.correct_answer:
                            question_score = question.points
                    elif question.question_type == 'multiple':
                        if set(user_answer) == set(question.correct_answer):
                            question_score = question.points
                
                score += question_score
                detailed_results.append({
                    'question_id': question.id,
                    'score': question_score,
                    'max_score': question.points
                })

            result = TestResult(
                assignment_id=assignment.id,
                test_id=test.id,
                user_id=session['user_id'],
                answers=answers,
                score=score,
                max_score=max_score,
                detailed_results=detailed_results
            )
            
            assignment.completed = True
            db.session.add(result)
            db.session.commit()

            return jsonify({
                'success': True,
                'result_id': result.id,
                'redirect_url': url_for('test_result', result_id=result.id)
            })

        except Exception as e:
            db.session.rollback()
            return jsonify({'error': str(e)}), 500

    return render_template('take_test.html', 
                         test=test,
                         assignment=assignment,
                         current_user={
                             'name': f"{user.surname} {user.name}",
                             'role': 'candidate'
                         })

@app.route('/test_result/<int:result_id>')
def test_result(result_id):
    if 'user_id' not in session or session.get('user_role') != 'candidate':
        flash('Доступ запрещен', 'error')
        return redirect(url_for('login'))
    
    user = db.session.get(User, session['user_id'])
    if not user:
        flash('Пользователь не найден', 'error')
        return redirect(url_for('login'))

    result = TestResult.query.get_or_404(result_id)
    if result.user_id != session['user_id']:
        flash('Доступ запрещен', 'error')
        return redirect(url_for('candidate'))
    
    return render_template('test_result.html',
                         result=result,
                         current_user={
                             'name': f"{user.surname} {user.name}",
                             'role': 'candidate'
                         })

@app.route('/upload_report', methods=['POST'])
def upload_report():
    if 'user_id' not in session or session.get('user_role') != 'hr':
        abort(403)
    
    file = request.files.get('report')
    if not file:
        abort(400, description="No file uploaded")
    
    hr = HR.query.filter_by(user_id=session['user_id']).first()
    if not hr:
        abort(403)
    
    file_type = file.mimetype
    allowed_types = [
        'application/pdf',
        'application/msword',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    ]
    
    if file_type not in allowed_types:
        abort(400, description="Unsupported file type")
    
    try:
        file_data = file.read()
        
        report = Report(
            hr_id=hr.id,
            name=request.form.get('name', file.filename),
            is_binary=True,
            file_type=file_type[:255],
            binary_data=file_data
        )
        
        db.session.add(report)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'report': {
                'id': report.id,
                'name': report.name,
                'created_at': report.created_at.isoformat()
            }
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/delete_report/<int:report_id>', methods=['POST'])
def delete_report(report_id):
    if 'user_id' not in session or session.get('user_role') != 'hr':
        abort(403)
    
    report = Report.query.get_or_404(report_id)
    hr = HR.query.filter_by(user_id=session['user_id']).first()
    
    if not hr or report.hr_id != hr.id:
        abort(403)
    
    try:
        db.session.delete(report)
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500
    
@app.route('/candidates/<int:candidate_id>/upload_resume', methods=['POST'])
def upload_resume(candidate_id):
    if 'user_id' not in session or session.get('user_role') != 'candidate':
        return jsonify({'error': 'Unauthorized'}), 401
    
    candidate = Candidate.query.get_or_404(candidate_id)
    if candidate.user_id != session['user_id']:
        return jsonify({'error': 'Forbidden'}), 403
    
    if 'resume' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['resume']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    try:
        candidate.resume_file = file.read()
        candidate.resume_filename = secure_filename(file.filename)
        candidate.resume_content_type = file.mimetype
        
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/upload_map', methods=['POST'])
def upload_map():
    if 'user_id' not in session or session.get('user_role') != 'hr':
        return jsonify({'error': 'Unauthorized'}), 401
    
    if 'name' not in request.form or 'image' not in request.files:
        return jsonify({'error': 'Missing required fields'}), 400
    
    name = request.form['name']
    image = request.files['image']
    
    if not name or not image or image.filename == '':
        return jsonify({'error': 'Invalid data'}), 400
    
    try:
        image_data = "data:" + image.mimetype + ";base64," + base64.b64encode(image.read()).decode('utf-8')
        
        hr = HR.query.filter_by(user_id=session['user_id']).first()
        new_map = OfficeMap(
            hr_id=hr.id,
            name=name,
            image_data=image_data
        )
        db.session.add(new_map)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'map_id': new_map.id
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500
    
@app.route('/download_resume/<int:candidate_id>')
def download_resume(candidate_id):
    if 'user_id' not in session or session.get('user_role') != 'candidate':
        abort(403)
    
    candidate = Candidate.query.get_or_404(candidate_id)
    if candidate.user_id != session['user_id']:
        abort(403)
    
    if not candidate.resume_file:
        abort(404)
    
    return Response(
        candidate.resume_file,
        mimetype=candidate.resume_content_type,
        headers={
            'Content-Disposition': f'attachment; filename="{candidate.resume_filename}"'
        }
    )

@app.route('/candidate')
def candidate():
    if 'user_id' not in session or session.get('user_role') != 'candidate':
        flash('Доступ запрещен', 'error')
        return redirect(url_for('login'))
    
    user = db.session.get(User, session['user_id'])
    if not user:
        flash('Пользователь не найден', 'error')
        return redirect(url_for('login'))
    
    candidate_record = Candidate.query.filter_by(user_id=user.id).first()
    if not candidate_record:
        flash('Запись кандидата не найдена', 'error')
        return redirect(url_for('login'))
    
    assigned_tests = TestAssignment.query.filter_by(
        candidate_id=candidate_record.id
    ).options(
        db.joinedload(TestAssignment.test),
        db.joinedload(TestAssignment.results)
    ).all()

    hr_contact = None
    if candidate_record.hr_id:
        hr_user = User.query.get(candidate_record.hr.user_id)
        if hr_user:
            hr_contact = {
                'name': f"{hr_user.surname} {hr_user.name}",
                'email': hr_user.mail,
                'phone': hr_user.phone
            }

    return render_template('candidate.html',
                         current_user=user,
                         hr_contact=hr_contact,
                         assigned_tests=assigned_tests,
                         current_year=datetime.now().year)

@app.route('/hr/save_note', methods=['POST'])
def save_hr_note():
    if 'user_id' not in session or session.get('user_role') != 'hr':
        return jsonify({'error': 'Unauthorized'}), 401
    
    hr_record = HR.query.filter_by(user_id=session['user_id']).first()
    if not hr_record:
        return jsonify({'error': 'HR record not found'}), 404
    
    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data provided'}), 400
    
    try:
        hr_record.note_title = data.get('title', '')
        hr_record.note_content = data.get('content', '')
        db.session.commit()
        
        return jsonify({
            'success': True,
            'note': {
                'title': hr_record.note_title,
                'content': hr_record.note_content
            }
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/hr/delete_note', methods=['POST'])
def delete_hr_note():
    if 'user_id' not in session or session.get('user_role') != 'hr':
        return jsonify({'error': 'Unauthorized'}), 401
    
    hr_record = HR.query.filter_by(user_id=session['user_id']).first()
    if not hr_record:
        return jsonify({'error': 'HR record not found'}), 404
    
    try:
        hr_record.note_title = None
        hr_record.note_content = None
        db.session.commit()
        
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/hr/get_note')
def get_hr_note():
    if 'user_id' not in session or session.get('user_role') != 'hr':
        return jsonify({'error': 'Unauthorized'}), 401
    
    hr_record = HR.query.filter_by(user_id=session['user_id']).first()
    if not hr_record:
        return jsonify({'error': 'HR record not found'}), 404
    
    if hr_record.note_title or hr_record.note_content:
        return jsonify({
            'note': {
                'title': hr_record.note_title,
                'content': hr_record.note_content
            }
        })
    else:
        return jsonify({'note': None})

@app.route('/hr/notes', methods=['GET', 'POST'])
def hr_notes():
    if 'user_id' not in session or session.get('user_role') != 'hr':
        return jsonify({'error': 'Unauthorized'}), 401
    
    hr = HR.query.filter_by(user_id=session['user_id']).first()
    if not hr:
        return jsonify({'error': 'HR not found'}), 404
    
    if request.method == 'GET':
        notes = [{
            'id': note.id,
            'title': note.title,
            'content': note.content,
            'created_at': note.created_at.isoformat()
        } for note in hr.notes]
        return jsonify({'notes': notes})
    
    elif request.method == 'POST':
        data = request.get_json()
        if not data or not data.get('title') or not data.get('content'):
            return jsonify({'error': 'Title and content are required'}), 400
        
        try:
            new_note = HRNote(
                hr_id=hr.id,
                title=data['title'],
                content=data['content']
            )
            db.session.add(new_note)
            db.session.commit()
            
            return jsonify({
                'success': True,
                'note': {
                    'id': new_note.id,
                    'title': new_note.title,
                    'content': new_note.content,
                    'created_at': new_note.created_at.isoformat()
                }
            })
        except Exception as e:
            db.session.rollback()
            return jsonify({'error': str(e)}), 500

@app.route('/hr/notes/<int:note_id>', methods=['PUT', 'DELETE'])
def hr_note(note_id):
    if 'user_id' not in session or session.get('user_role') != 'hr':
        return jsonify({'error': 'Unauthorized'}), 401
    
    hr = HR.query.filter_by(user_id=session['user_id']).first()
    if not hr:
        return jsonify({'error': 'HR not found'}), 404
    
    note = HRNote.query.filter_by(id=note_id, hr_id=hr.id).first()
    if not note:
        return jsonify({'error': 'Note not found'}), 404
    
    if request.method == 'PUT':
        data = request.get_json()
        if not data or not data.get('title') or not data.get('content'):
            return jsonify({'error': 'Title and content are required'}), 400
        
        try:
            note.title = data['title']
            note.content = data['content']
            db.session.commit()
            return jsonify({
                'success': True,
                'note': {
                    'id': note.id,
                    'title': note.title,
                    'content': note.content
                }
            })
        except Exception as e:
            db.session.rollback()
            return jsonify({'error': str(e)}), 500
    
    elif request.method == 'DELETE':
        try:
            db.session.delete(note)
            db.session.commit()
            return jsonify({'success': True})
        except Exception as e:
            db.session.rollback()
            return jsonify({'error': str(e)}), 500

@app.route('/hr/compare_candidates_data')
def compare_candidates_data():
    if 'user_id' not in session or session.get('user_role') != 'hr':
        return jsonify({'error': 'Unauthorized'}), 401
    
    candidate_ids = request.args.get('candidate_ids', '').split(',')
    if not candidate_ids or len(candidate_ids) < 2:
        return jsonify({'error': 'Необходимо выбрать минимум 2 кандидата'}), 400
    
    try:
        hr = HR.query.filter_by(user_id=session['user_id']).first()
        resume_test = Test.query.filter_by(hr_id=hr.id, name="Резюме").first()
        
        if not resume_test:
            return jsonify({'error': 'Тест "Резюме" не найден'}), 404
        
        candidates_data = []
        questions_data = []
        
        questions = TestQuestion.query.filter_by(test_id=resume_test.id).order_by(TestQuestion.order_num).all()
        for question in questions:
            questions_data.append({
                'id': question.id,
                'text': question.question_text,
                'type': question.question_type,
                'correct_answer': question.correct_answer,
                'points': question.points
            })
        
        for candidate_id in candidate_ids:
            result = TestResult.query.filter_by(
                test_id=resume_test.id,
                user_id=candidate_id
            ).first()
            
            if not result:
                continue
            
            candidate = Candidate.query.filter_by(user_id=candidate_id).first()
            user = User.query.get(candidate_id)
            
            candidates_data.append({
                'id': candidate_id,
                'name': f"{user.surname} {user.name}",
                'vacancy': candidate.status,
                'answers': result.answers,
                'question_scores': {dr['question_id']: dr['score'] for dr in result.detailed_results}
            })
        
        return jsonify({
            'questions': questions_data,
            'candidates': candidates_data
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500
             
@app.route('/hr/compare_candidates', methods=['GET', 'POST'])
def compare_candidates():
    if 'user_id' not in session or session.get('user_role') != 'hr':
        flash('Доступ запрещен', 'error')
        return redirect(url_for('login'))
    
    hr = HR.query.filter_by(user_id=session['user_id']).first()
    if not hr:
        flash('HR запись не найдена', 'error')
        return redirect(url_for('login'))
    
    resume_test = Test.query.filter_by(hr_id=hr.id, name="Резюме").first()
    
    if not resume_test:
        flash('Тест "Резюме" не найден', 'error')
        return redirect(url_for('hr'))
    
    if request.method == 'POST':
        selected_answers = request.get_json().get('selected_answers', [])
        
        candidate_scores = {}
        for answer in selected_answers:
            candidate_id = answer['candidate_id']
            points = answer['points']
            candidate_scores[candidate_id] = candidate_scores.get(candidate_id, 0) + points
        
        max_score = max(candidate_scores.values()) if candidate_scores else 0
        
        best_candidates = [cand_id for cand_id, score in candidate_scores.items() if score == max_score]
        
        return jsonify({
            'best_candidates': best_candidates,
            'scores': candidate_scores
        })
    
    candidates = Candidate.query.filter_by(hr_id=hr.id).options(
        db.joinedload(Candidate.user),
        db.joinedload(Candidate.test_assignments).joinedload(TestAssignment.results)
    ).all()
    
    candidates_with_resume = []
    for candidate in candidates:
        assignment = next(
            (a for a in candidate.test_assignments 
             if a.test_id == resume_test.id and a.completed and a.results),
            None
        )
        if assignment:
            candidates_with_resume.append({
                'id': candidate.id,
                'user_id': candidate.user_id,
                'name': f"{candidate.user.surname} {candidate.user.name}",
                'vacancy': candidate.status,
                'results': assignment.results
            })
    
    return render_template('compare_candidates.html',
                         candidates=candidates_with_resume,
                         resume_test=resume_test,
                         current_user={
                             'name': session.get('user_name')
                         })

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        user_type = request.form.get('user_type')
        surname = request.form.get('surname')
        name = request.form.get('name')
        patronymic = request.form.get('patronymic', '')
        email = request.form.get('email')
        phone = request.form.get('phone')
        password = request.form.get('password')
        photo = request.files.get('photo')
        vacancy = request.form.get('vacancy', '')
        
        errors = []
        
        if not all([surname, name, email, phone, password]):
            errors.append('Все обязательные поля должны быть заполнены')
        if not is_valid_email(email):
            errors.append('Некорректный формат email')
        if not is_valid_phone(phone):
            errors.append('Некорректный номер телефона')
        if len(password) < 6:
            errors.append('Пароль должен содержать минимум 6 символов')
        if User.query.filter_by(mail=email).first():
            errors.append('Пользователь с таким email уже существует')

        if errors:
            for error in errors:
                flash(error, 'error')
            return redirect(url_for('register'))

        photo_filename = save_uploaded_file(photo)

        try:
            new_user = User(
                role=user_type,
                surname=surname,
                name=name,
                patronymic=patronymic,
                mail=email,
                phone=phone,
                password=password,
                image=photo_filename
            )
            db.session.add(new_user)
            db.session.flush()

            if user_type == 'hr':
                new_hr = HR(
                    user_id=new_user.id,
                    department='Default Department'
                )
                db.session.add(new_hr)

            elif user_type == 'candidate':
                hr_phone = request.form.get('hr_phone')
                hr_user = User.query.filter_by(phone=hr_phone, role='hr').first()
                
                if not hr_user:
                    flash('HR с указанным номером телефона не найден', 'error')
                    return redirect(url_for('register'))
                
                hr_record = HR.query.filter_by(user_id=hr_user.id).first()
                if not hr_record:
                    hr_record = HR(
                    user_id=hr_user.id,
                    department='Default Department'
                    )
                    db.session.add(hr_record)
                    db.session.flush()
                
                candidate = Candidate(
                    user_id=new_user.id,
                    hr_id=hr_record.id,
                    status=vacancy
                )
                db.session.add(candidate)

            db.session.commit()
            flash('Регистрация прошла успешно! Теперь вы можете войти.', 'success')
            return redirect(url_for('entry'))
            
        except Exception as e:
            db.session.rollback()
            flash(f'Ошибка при регистрации: {str(e)}', 'error')
            app.logger.error(f'Registration error: {str(e)}')

    hr_users = User.query.filter_by(role='hr').all()
    valid_hrs = [u for u in hr_users if HR.query.filter_by(user_id=u.id).first()]
    return render_template('register.html', hr_users=valid_hrs)

@app.route('/hr/report_editor')
def report_editor():
    if 'user_id' not in session or session.get('user_role') != 'hr':
        flash('Доступ запрещен', 'error')
        return redirect(url_for('login'))
    return render_template('report_editor.html')

@app.route('/reports/<int:report_id>')
def view_report(report_id):
    if 'user_id' not in session or session.get('user_role') != 'hr':
        flash('Доступ запрещен', 'error')
        return redirect(url_for('login'))
    
    report = Report.query.get_or_404(report_id)
    if report.hr.user_id != session['user_id']:
        flash('Нет прав доступа к этому отчету', 'error')
        return redirect(url_for('hr'))
    
    return render_template('view_report.html', 
                         report=report,
                         current_user={
                             'name': session.get('user_name')
                         })

@app.route('/reports/<int:report_id>/download')
def download_report(report_id):
    if 'user_id' not in session or session.get('user_role') != 'hr':
        abort(403)
    
    report = Report.query.get_or_404(report_id)
    hr = HR.query.filter_by(user_id=session['user_id']).first()
    
    if not hr or report.hr_id != hr.id:
        abort(403)
    
    if report.is_binary and report.binary_data:
        ext = 'bin'
        if 'pdf' in report.file_type:
            ext = 'pdf'
        elif 'word' in report.file_type or 'msword' in report.file_type:
            ext = 'docx' if 'openxml' in report.file_type else 'doc'
        
        return Response(
            report.binary_data,
            mimetype=report.file_type,
            headers={'Content-Disposition': f'attachment; filename={report.name}.{ext}'}
        )
    else:
        return Response(
            report.text,
            mimetype='text/plain',
            headers={'Content-Disposition': f'attachment; filename={report.name}.txt'}
        )

@app.route('/reports/<int:report_id>/export/<format>')
def export_report(report_id, format):
    if 'user_id' not in session or session.get('user_role') != 'hr':
        abort(403)
    
    report = Report.query.get_or_404(report_id)
    hr = HR.query.filter_by(user_id=session['user_id']).first()
    
    if not hr or report.hr_id != hr.id:
        abort(403)
    
    if format == 'pdf':
        from io import BytesIO
        from reportlab.pdfgen import canvas
        
        buffer = BytesIO()
        p = canvas.Canvas(buffer)
        p.drawString(100, 100, report.name)
        p.drawString(100, 120, report.text[:500] if report.text else "Бинарный файл - скачайте оригинал")
        p.showPage()
        p.save()
        buffer.seek(0)
        
        return Response(
            buffer,
            mimetype='application/pdf',
            headers={'Content-Disposition': f'attachment; filename={report.name}.pdf'}
        )
    
    elif format == 'word':
        from docx import Document
        from io import BytesIO
        
        document = Document()
        document.add_heading(report.name, 0)
        document.add_paragraph(report.text if report.text else "Бинарный файл - скачайте оригинал")
        
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        
        return Response(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment; filename={report.name}.docx'}
        )
    
    else:
        abort(400, description="Unsupported format")

@app.route('/api/hr_candidates')
def get_hr_candidates():
    if 'user_id' not in session or session.get('user_role') != 'hr':
        return jsonify({'error': 'Unauthorized'}), 401
    
    hr = HR.query.filter_by(user_id=session['user_id']).first()
    if not hr:
        return jsonify({'error': 'HR not found'}), 404
    
    candidates = Candidate.query.filter_by(hr_id=hr.id).options(
        db.joinedload(Candidate.user)
    ).all()
    
    return jsonify([
        {
            'id': cand.id,
            'name': f"{cand.user.surname} {cand.user.name}",
            'email': cand.user.mail,
            'phone': cand.user.phone,
            'status': cand.status
        } for cand in candidates
    ])

@app.route('/api/save_report', methods=['POST'])
def save_report():
    if 'user_id' not in session or session.get('user_role') != 'hr':
        return jsonify({'error': 'Unauthorized'}), 401
    
    data = request.get_json()
    hr = HR.query.filter_by(user_id=session['user_id']).first()
    
    try:
        new_report = Report(
            hr_id=hr.id,
            name=data['title'],
            text=data['content']
        )
        db.session.add(new_report)
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/hr/delete_test/<int:test_id>', methods=['POST'])
def delete_test(test_id):
    if 'user_id' not in session or session.get('user_role') != 'hr':
        return jsonify({'error': 'Unauthorized'}), 401
    
    try:
        test = Test.query.get_or_404(test_id)
        db.session.delete(test)
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500
    
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')
        
        user = User.query.filter_by(mail=email).first()
        
        if user and check_password_hash(user.password, password):
            session['user_id'] = user.id
            session['user_role'] = user.role
            session['user_name'] = f"{user.surname} {user.name}"
            session['user_fullname'] = f"{user.surname} {user.name}"
            
            flash(f'Добро пожаловать, {user.surname}!', 'success')
            
            if user.role == 'admin':
                return redirect(url_for('admin'))
            elif user.role == 'hr':
                return redirect(url_for('hr'))
            elif user.role == 'employee':
                return redirect(url_for('employee'))
            elif user.role == 'candidate':
                return redirect(url_for('candidate'))
            else:
                return redirect(url_for('entry'))
        else:
            flash('Неверный email или пароль', 'error')
    
    return render_template('entry.html')

@app.route('/logout')
def logout():
    session.clear()
    flash('Вы успешно вышли из системы', 'success')
    return redirect(url_for('entry'))